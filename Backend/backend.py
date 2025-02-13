import cv2
import pytesseract
import time
from langdetect import detect
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
from docx import Document
from docx.shared import Inches
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from scipy.spatial.distance import cosine
from groq import Groq
import io
from PIL import Image
from flask import Flask, request, send_file, jsonify, after_this_request
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
OUTPUT_WORD_FILE = "lecture_analysis.docx"
IMAGE_OUTPUT_DIR = "extracted_images"
MAX_WORKERS = 4
TEXT_SIMILARITY_THRESHOLD = 0.7
IMAGE_SIMILARITY_THRESHOLD = 0.95

def perform_ocr(frame, timestamp):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray)
    return text, timestamp

def filter_text(text):
    try:
        lang = detect(text)
        return lang == 'en' and len(text.strip()) > 0
    except:
        return False

def detect_and_save_images(frame, timestamp):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    cnts = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]

    image_filenames = []
    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        if w > 100 and h > 100:
            roi = frame[y:y+h, x:x+w]
            if is_likely_image(roi):
                if not os.path.exists(IMAGE_OUTPUT_DIR):
                    os.makedirs(IMAGE_OUTPUT_DIR)
                filename = f"{IMAGE_OUTPUT_DIR}/image_{timestamp.replace(':', '-')}_{len(image_filenames)}.png"

                # --- ADDED ERROR HANDLING ---
                try:
                    success = cv2.imwrite(filename, roi)
                    if success:
                        image_filenames.append(filename)
                    else:
                        print(f"ERROR: Failed to write image file: {filename}")
                except Exception as e:
                    print(f"ERROR: Exception while writing image file: {filename} - {e}")
                # --- END ADDED ERROR HANDLING ---

    return image_filenames

def is_likely_image(roi):
    gray_roi = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    std_dev = np.std(gray_roi)
    return std_dev > 20

def process_frame(frame, current_time):
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(current_time))
    text, _ = perform_ocr(frame, timestamp)
    image_filenames = []
    if filter_text(text):
        image_filenames = detect_and_save_images(frame, timestamp)
    return text, timestamp, image_filenames

def text_similarity(text1, text2):
    vectorizer = TfidfVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    return 1 - cosine(vectors[0], vectors[1])

def image_similarity(img1_path, img2_path):
    img1 = cv2.imread(img1_path, cv2.IMREAD_GRAYSCALE)
    img2 = cv2.imread(img2_path, cv2.IMREAD_GRAYSCALE)
    if img1 is None or img2 is None:
        return 0.0

    img1 = cv2.resize(img1, (300, 300))
    img2 = cv2.resize(img2, (300, 300))
    similarity = cv2.matchTemplate(img1, img2, cv2.TM_CCOEFF_NORMED)
    return np.max(similarity)

def describe_image(image_path):
    try:
        img = cv2.imread(image_path)
        if img is None:
            return f"[Error: Could not load image at {image_path}]"

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        ocr_text = pytesseract.image_to_string(gray)
        height, width = img.shape[:2]
        brightness = np.mean(gray)
        contrast = np.std(gray)
        edges = cv2.Canny(gray, 100, 200)
        edge_density = np.sum(edges > 0) / (height * width)

        image_type = "text-heavy" if len(ocr_text.strip()) > 100 else \
                    "diagram" if edge_density > 0.1 else \
                    "photo"

        description = f"""Image Description:
        - Type: {image_type}
        - Dimensions: {width}x{height}
        - Content: {ocr_text.strip() if ocr_text.strip() else 'No text detected'}
        - Visual characteristics: {'High contrast with detailed edges' if edge_density > 0.1 else 'Smooth transitions with natural appearance'}
        """
        return description
    except Exception as e:
        print(f"Error describing image {image_path}: {str(e)}")
        return f"[Error analyzing image: {image_path}]"

def get_groq_explanation(extracted_content):
    try:
        print(f"DEBUG: extracted_content in get_groq_explanation: {extracted_content}")  # Add this

        formatted_content = ""
        for text, timestamp, image_filenames in extracted_content:
            formatted_content += f"\n=== Timestamp: {timestamp} ===\n"
            formatted_content += f"Text Content:\n{text}\n"

            if image_filenames:
                formatted_content += "\nVisual Content:\n"
                for img_file in image_filenames:
                    image_desc = describe_image(img_file)
                    formatted_content += f"{image_desc}\n"

            formatted_content += "\n" + "="*50 + "\n"

        client = Groq(api_key="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") # Replace with your Groq API key
        #TEMPORARY PROMPT
        prompt = """Analyze the following lecture content, including both text and visual elements. Please provide:

        1. Comprehensive Analysis:
           - Main concepts or each point explained in detail
           - Integration of visual elements with textual content
           - Technical terminology explained
        
        2. Visual Content Analysis:
           - How the images support the lecture content
           - Key visual elements and their significance
           - Relationship between text and visual aids
           -Also identify the different images in that image if any
        
        3. Summary and Practical Applications:
           - Key takeaways
           - Real-world applications
           - Learning objectives achieved
        5.Alternate explanation of the slides.
        4.If you find any numerical in the text provided by me also solve that numerical
        Lecture Content:
        5. If you find any educational diagram like nfa or dfa and you feel that its construction and working is important for the students to understand for the exam, please explain both with the help of a diagram.
        {formatted_content}

        Please structure your response with clear sections and subsections, relating the visual elements to the concepts discussed.""" + formatted_content[:500]
        # prompt = f"""Analyze the following lecture content... (rest of your prompt) ..."""  # Your Groq prompt

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama3-8b-8192",
        )
        print(f"DEBUG: Groq response: {chat_completion.choices[0].message.content}") # Add this
        return chat_completion.choices[0].message.content
    except Exception as e:
        print(f"Error getting Groq explanation: {str(e)}")
        return None

def create_enhanced_document(extracted_content, groq_explanation):
    doc = Document()
    doc.add_heading('Lecture Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading('Original Lecture Content', 1)
    for text, timestamp, image_filenames in extracted_content:
        doc.add_heading(f"Timestamp: {timestamp}", 2)
        doc.add_paragraph(text)
        if image_filenames:
            doc.add_heading("Visual Content", 3)
            for image_filename in image_filenames:
                try:
                    doc.add_picture(image_filename, width=Inches(6))
                    desc = describe_image(image_filename)
                    doc.add_paragraph(desc, style='Caption')
                except Exception as e:
                    print(f"Error adding image to document: {e}")
                    doc.add_paragraph(f"[Error: Could not add image {image_filename}]")
        doc.add_paragraph()
    print(f"DEBUG: groq_explanation in create_enhanced_document: {groq_explanation}")  # Add this
    if groq_explanation:
        doc.add_heading('Comprehensive Analysis', 1)
        doc.add_paragraph(groq_explanation)

    doc.add_heading('Technical Appendix', 1)
    doc.add_paragraph(f"""
    Processing Details:
    - Content Extraction Interval: Dynamically calculated
    - Number of Images Processed: {sum(len(img) for _, _, img in extracted_content)}
    - Analysis Model: llama3-8b-8192
    """)
    return doc


@app.route('/analyze_video', methods=['POST'])
def analyze_video():
    temp_video_path = None
    temp_docx_path = None  # Initialize temp_docx_path as well
    try:
        # ... (rest of your input handling and video loading) ...
        if 'video' not in request.files:
            return jsonify({'error': 'No video file provided'}), 400
        if not request.form.get('start_phrase') or not request.form.get('end_phrase'):
            return jsonify({'error': 'start_phrase and end_phrase are required'}), 400

        video_file = request.files['video']
        start_phrase = request.form['start_phrase']
        end_phrase = request.form['end_phrase']

        if video_file.filename == '':
            return jsonify({'error': 'No video file selected'}), 400

        temp_video_path = "temp_video.mp4"
        video_file.save(temp_video_path)

        cap = cv2.VideoCapture(temp_video_path)
        if not cap.isOpened():
            return jsonify({'error': 'Could not open video file'}), 500
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        video_duration_seconds = frame_count / fps
        dynamic_interval_seconds = min(video_duration_seconds / 100, 3 * 60)
        frame_interval = int(fps * dynamic_interval_seconds)

        first_phrase_found = False
        second_phrase_found = False
        extracted_content = []
        current_text = ""
        saved_images = []
        start_timestamp = None
        end_timestamp = None

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = []
            for frame_number in range(0, frame_count, frame_interval):
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
                ret, frame = cap.read()
                if not ret:
                    break
                current_time = frame_number / fps
                futures.append(executor.submit(process_frame, frame, current_time))

            for future in as_completed(futures):
                text, timestamp, image_filenames = future.result()
                if text:
                    if not first_phrase_found:
                        if start_phrase in text:
                            first_phrase_found = True
                            start_timestamp = timestamp
                            current_text = text
                            extracted_content.append((text, timestamp, image_filenames))
                            saved_images.extend(image_filenames)
                    elif end_phrase in text:
                        second_phrase_found = True
                        end_timestamp = timestamp
                        extracted_content.append((text, timestamp, image_filenames))
                        saved_images.extend(image_filenames)
                        break
                    elif first_phrase_found and text_similarity(current_text, text) < TEXT_SIMILARITY_THRESHOLD:
                        current_text = text
                        new_images = []
                        for img in image_filenames:
                            if not any(image_similarity(img, saved_img) > IMAGE_SIMILARITY_THRESHOLD for saved_img in saved_images):
                                new_images.append(img)
                        if new_images:
                            extracted_content.append((text, timestamp, new_images))
                            saved_images.extend(new_images)
        cap.release()

        if first_phrase_found and second_phrase_found:
            filtered_content = []
            for text, timestamp, image_filenames in extracted_content:
                if start_timestamp <= timestamp <= end_timestamp:
                    filtered_content.append((text, timestamp, image_filenames))
            print(f"DEBUG: filtered_content in analyze_video: {filtered_content}")
            groq_explanation = get_groq_explanation(filtered_content)
            doc = create_enhanced_document(filtered_content, groq_explanation)
            temp_docx_path = "temp_lecture_analysis.docx"
            doc.save(temp_docx_path)

            # --- KEY CHANGE: Explicitly close the doc object ---
            doc = None  # Release the object
            # -------------

            all_saved_images = set()
            for _, _, image_filenames_list in filtered_content:
                all_saved_images.update(image_filenames_list)

            for filename in os.listdir(IMAGE_OUTPUT_DIR):
                file_path = os.path.join(IMAGE_OUTPUT_DIR, filename)
                if file_path not in all_saved_images:
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"Error removing unused image {file_path}: {str(e)}")

            # --- Modified after_this_request ---
            @after_this_request
            def remove_files(response):
                attempts = 0
                max_attempts = 5  # Increased attempts
                delay = 1.0      # Increased delay
                while attempts < max_attempts:
                    try:
                        if temp_docx_path and os.path.exists(temp_docx_path):
                            os.remove(temp_docx_path)
                        if temp_video_path and os.path.exists(temp_video_path):
                            os.remove(temp_video_path)
                        return response  # Files deleted successfully
                    except PermissionError:
                        attempts += 1
                        print(f"File deletion failed (attempt {attempts}). Retrying in {delay} seconds...")
                        time.sleep(delay)
                # If we reach here, deletion failed after all attempts
                app.logger.error(f"Failed to delete temporary files after multiple attempts.")
                return response
            return send_file(temp_docx_path, as_attachment=True)

        else:
            return jsonify({'error': 'Start or end phrase not found in the video'}), 404

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if temp_video_path and os.path.exists(temp_video_path):
            try:
                os.remove(temp_video_path)
            except Exception as e:
                app.logger.error(f"Error removing temp_video_path in finally block: {e}")
        # ---Crucially, also attempt to delete temp_docx_path in finally---
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception as e:
                app.logger.error(f"Error removing temp_docx_path in finally block: {e}")

if __name__ == '__main__':
    app.run(debug=True)