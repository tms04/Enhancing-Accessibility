import cv2
import pytesseract
import time
from langdetect import detect
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
from docx import Document
from docx.shared import Inches
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from scipy.spatial.distance import cosine
from PIL import Image
from flask import Flask, request, send_file, jsonify, after_this_request
from flask_cors import CORS
import io
import base64
import google.generativeai as genai
from config import GEMINI_API_KEY, GENAI_MODEL_NAME

app = Flask(__name__)
CORS(app)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Or your path
OUTPUT_WORD_FILE = "lecture_analysis_gemini.docx"
IMAGE_OUTPUT_DIR = "extracted_images"
MAX_WORKERS = 4
TEXT_SIMILARITY_THRESHOLD = 0.7
IMAGE_SIMILARITY_THRESHOLD = 0.95
MIN_CHANGE_THRESHOLD = 0.2  # Threshold for considering a change (adjust as needed)

try:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel(GENAI_MODEL_NAME)
except Exception as e:
    print(f"Error configuring Gemini API: {e}")
    gemini_model = None

def preprocess_image_for_handwriting(image_path):
    try:
        img = cv2.imread(image_path)
        if img is None:
            return None

        preprocessed_versions = []

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh1 = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        kernel = np.ones((2,2), np.uint8)
        eroded1 = cv2.erode(thresh1, kernel, iterations=1)
        path1 = f"{image_path}_prep_v1.png"
        cv2.imwrite(path1, eroded1)
        preprocessed_versions.append(path1)

        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        gray2 = clahe.apply(gray)
        thresh2 = cv2.threshold(gray2, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        path2 = f"{image_path}_prep_v2.png"
        cv2.imwrite(path2, thresh2)
        preprocessed_versions.append(path2)

        gray3 = cv2.GaussianBlur(gray, (3, 3), 0)
        edges = cv2.Canny(gray3, 50, 150)
        dilated_edges = cv2.dilate(edges, kernel, iterations=1)
        path3 = f"{image_path}_prep_v3.png"
        cv2.imwrite(path3, dilated_edges)
        preprocessed_versions.append(path3)

        return preprocessed_versions
    except Exception as e:
        print(f"Error preprocessing image: {str(e)}")
        return None

def extract_handwritten_text(image_path):
    try:
        preprocessed_paths = preprocess_image_for_handwriting(image_path)
        if not preprocessed_paths:
            return None

        results = []

        for i, prep_path in enumerate(preprocessed_paths):
            if i == 0:
                config = r'--oem 3 --psm 6 -l eng+handwritten'
            elif i == 1:
                config = r'--oem 3 --psm 6 -l eng+equ'
            else:
                config = r'--oem 3 --psm 4 -l eng+equ+handwritten'

            text = pytesseract.image_to_string(
                Image.open(prep_path),
                config=config
            )
            results.append(text.strip())

        for path in preprocessed_paths:
            if os.path.exists(path):
                os.remove(path)

        combined_text = post_process_extracted_text(results)
        return combined_text
    except Exception as e:
        print(f"Error extracting handwritten text: {str(e)}")
        return None

def post_process_extracted_text(text_versions):
    if not text_versions:
        return None

    longest_text = max(text_versions, key=len)
    final_text = longest_text

    symbol_corrections = {
        '~': '∪',
        '\\': '∩',
        '>': '⟩',
        '<': '⟨',
        '2': 'λ',
        '->': '→',
        '=>': '⇒',
        '--': '⟶',
        '(': '{',
        ')': '}',
        '|': '∈',
        '#': '≠',
        '?': 'λ',
    }

    for incorrect, correct in symbol_corrections.items():
        if incorrect in final_text and ('=' in final_text or '{' in final_text or '}' in final_text):
            final_text = final_text.replace(incorrect, correct)

    if 'sigma' in final_text.lower() or 'sigma*' in final_text.lower():
        final_text = final_text.replace('sigma', 'Σ').replace('Sigma', 'Σ')

    if ('lambda' in final_text.lower() or 'empty' in final_text.lower() or
            '4' in final_text and '{4' in final_text):
        final_text = final_text.replace('lambda', 'λ').replace('Lambda', 'λ').replace('{4', '{λ')

    return final_text

def perform_ocr(frame, timestamp):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray)
    return text, timestamp

def filter_text(text):
    try:
        lang = detect(text)
        return lang == 'en' and len(text.strip()) > 0
    except:
        return False

def detect_and_save_images(frame, timestamp):
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    cnts = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]

    image_filenames = []
    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        if w > 100 and h > 100:
            roi = frame[y:y+h, x:x+w]
            if is_likely_image(roi):
                if not os.path.exists(IMAGE_OUTPUT_DIR):
                    os.makedirs(IMAGE_OUTPUT_DIR)
                filename = f"{IMAGE_OUTPUT_DIR}/image_{timestamp.replace(':', '-')}_{len(image_filenames)}.png"
                try:
                    success = cv2.imwrite(filename, roi)
                    if success:
                        image_filenames.append(filename)
                    else:
                        print(f"ERROR: Failed to write image file: {filename}")
                except Exception as e:
                    print(f"ERROR: Exception while writing image file: {filename} - {e}")
    return image_filenames

def is_likely_image(roi):
    gray_roi = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    std_dev = np.std(gray_roi)
    return std_dev > 20

def process_frame(frame, current_time):
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(current_time))
    text, _ = perform_ocr(frame, timestamp)
    image_filenames = []
    if filter_text(text):
        image_filenames = detect_and_save_images(frame, timestamp)
    return text, timestamp, image_filenames, frame

def text_similarity(text1, text2):
    vectorizer = TfidfVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    return 1 - cosine(vectors[0], vectors[1])

def image_similarity(img1_path, img2_path):
    img1 = cv2.imread(img1_path, cv2.IMREAD_GRAYSCALE)
    img2 = cv2.imread(img2_path, cv2.IMREAD_GRAYSCALE)
    if img1 is None or img2 is None:
        return 0.0

    img1 = cv2.resize(img1, (300, 300))
    img2 = cv2.resize(img2, (300, 300))
    similarity = cv2.matchTemplate(img1, img2, cv2.TM_CCOEFF_NORMED)
    return np.max(similarity)

def calculate_image_change(prev_frame, current_frame):
    if prev_frame is None or current_frame is None:
        return 1.0

    prev_gray = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
    current_gray = cv2.cvtColor(current_frame, cv2.COLOR_BGR2GRAY)

    if prev_gray.shape != current_gray.shape:
        current_resized = cv2.resize(current_gray, (prev_gray.shape[1], prev_gray.shape[0]))
        diff = cv2.absdiff(prev_gray, current_resized)
    else:
        diff = cv2.absdiff(prev_gray, current_gray)

    normalized_diff = np.sum(diff) / (prev_gray.shape[0] * prev_gray.shape[1] * 255)
    return normalized_diff

def describe_image(image_path):
    try:
        img = cv2.imread(image_path)
        if img is None:
            return f"[Error: Could not load image at {image_path}]"

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        ocr_text = pytesseract.image_to_string(gray)

        math_symbols = ['Σ', 'λ', '∪', '∩', '∈', '→', '⇒', '{', '}', '*', '=']
        math_words = ['set', 'function', 'operation', 'theorem', 'proof', 'algorithm', 'lambda', 'sigma']

        is_math_content = False
        for symbol in math_symbols:
            if symbol in ocr_text:
                is_math_content = True
                break

        if not is_math_content:
            for word in math_words:
                if word.lower() in ocr_text.lower():
                    is_math_content = True
                    break

        height, width = img.shape[:2]
        brightness = np.mean(gray)
        contrast = np.std(gray)
        edges = cv2.Canny(gray, 100, 200)
        edge_density = np.sum(edges > 0) / (height * width)

        if is_math_content:
            image_type = "mathematical notation"
        elif len(ocr_text.strip()) > 100:
            image_type = "text-heavy"
        elif edge_density > 0.1:
            image_type = "diagram"
        else:
            image_type = "photo"

        description = f"""Image Description:
        - Type: {image_type}
        - Dimensions: {width}x{height}
        - Content: {ocr_text.strip() if ocr_text.strip() else 'No text detected'}
        - Visual characteristics: {'High contrast with detailed edges' if edge_density > 0.1 else 'Smooth transitions with natural appearance'}
        """

        if is_math_content:
            description += "\n- Contains mathematical notation or formulas"

        return description
    except Exception as e:
        print(f"Error describing image {image_path}: {str(e)}")
        return f"[Error analyzing image: {image_path}]"

def get_gemini_explanation(extracted_content):
    try:
        print("DEBUG: Inside get_gemini_explanation")
        contents = []

        for text, timestamp, image_filenames in extracted_content:
            text_part = f"=== Timestamp: {timestamp} ===\nExtracted Text:\n{text}\n"
            contents.append(text_part)

            for img_file in image_filenames:
                try:
                    with open(img_file, "rb") as image_file:
                        image_data = image_file.read()
                        base64_image = base64.b64encode(image_data).decode("utf-8")
                        mime_type = "image/png"
                        image_part = {
                            "mime_type": mime_type,
                            "data": base64_image
                        }
                        contents.append(image_part)

                except Exception as e:
                    print(f"Error reading image {img_file}: {e}")

        prompt_text = """Analyze the lecture content, including text and images.
        Note that the OCR-extracted text may not include handwritten text, so you must carefully examine the images for any handwritten notes or mathematical expressions.
        Explain the key concepts, interpret any handwritten text or mathematical notation from the images,
        and describe how the images relate to the text. Provide a summary and any real-world applications."""

        if gemini_model:
            if GENAI_MODEL_NAME == "gemini-1.5-flash":
                gemini_response = gemini_model.generate_content(
                    contents=[prompt_text] + contents,
                )
            else:
                gemini_response = gemini_model.generate_content(
                    prompt_text + "\n" + str(contents)
                )

            if gemini_response.text:
                return gemini_response.text
            else:
                return "No explanation available from Gemini."
        else:
            return "Gemini API is not configured. Please check your API key and network connection."

    except Exception as e:
        print(f"Error getting Gemini explanation: {str(e)}")
        return f"Error: {str(e)}"

def create_enhanced_document(extracted_content, gemini_explanation):
    doc = Document()
    doc.add_heading('Lecture Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    math_note = doc.add_paragraph()
    math_note.add_run("Note: Mathematical symbols and notation are best viewed in the original images.").bold = True

    doc.add_heading('Original Lecture Content', 1)
    for text, timestamp, image_filenames in extracted_content:
        doc.add_heading(f"Timestamp: {timestamp}", 2)
        doc.add_paragraph(text)
        if image_filenames:
            doc.add_heading("Visual Content", 3)
            for image_filename in image_filenames:
                try:
                    doc.add_picture(image_filename, width=Inches(6))
                    handwritten_text = extract_handwritten_text(image_filename)
                    if handwritten_text:
                        p = doc.add_paragraph("Extracted Text: ")
                        p.add_run(handwritten_text)
                    desc = describe_image(image_filename)
                    doc.add_paragraph(desc, style='Caption')
                except Exception as e:
                    print(f"Error adding image to document: {e}")
                    doc.add_paragraph(f"[Error: Could not add image {image_filename}]")
        doc.add_paragraph()

    if gemini_explanation:
        doc.add_heading('Comprehensive Analysis', 1)
        doc.add_paragraph(gemini_explanation)

    doc.add_heading('Technical Appendix', 1)
    doc.add_paragraph(f"""
    Processing Details:
    - Content Extraction Interval: Dynamically calculated
    - Number of Images Processed: {sum(len(img) for _, _, img in extracted_content)}
    - Analysis Model: {GENAI_MODEL_NAME}
    - Text Extraction: Enhanced OCR with mathematical symbol detection
    """)
    return doc

@app.route('/analyze_video', methods=['POST'])
def analyze_video():
    temp_video_path = None

    try:
        if 'video' not in request.files:
            return jsonify({'error': 'No video file provided'}), 400
        if not request.form.get('start_phrase') or not request.form.get('end_phrase'):
            return jsonify({'error': 'start_phrase and end_phrase are required'}), 400

        video_file = request.files['video']
        start_phrase = request.form['start_phrase']
        end_phrase = request.form['end_phrase']

        if video_file.filename == '':
            return jsonify({'error': 'No video file selected'}), 400

        temp_video_path = "temp_video.mp4"
        video_file.save(temp_video_path)

        cap = cv2.VideoCapture(temp_video_path)
        if not cap.isOpened():
            return jsonify({'error': 'Could not open video file'}), 500

        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        video_duration_seconds = frame_count / fps

        initial_interval_seconds = min(3 * 60, video_duration_seconds / 100)
        initial_frame_interval = int(fps * initial_interval_seconds)
        post_start_interval_seconds = initial_interval_seconds / 3
        post_start_frame_interval = int(fps * post_start_interval_seconds)

        first_phrase_found = False
        second_phrase_found = False
        extracted_content = []
        current_text = ""
        saved_images = []
        start_timestamp = None
        end_timestamp = None
        previous_frame = None
        previous_text = None
        previous_images = []

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = []
            current_frame_interval = initial_frame_interval

            for frame_number in range(0, frame_count, current_frame_interval):
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
                ret, current_frame = cap.read()
                if not ret:
                    break
                current_time = frame_number / fps
                futures.append(executor.submit(process_frame, current_frame.copy(), current_time))

                if first_phrase_found and current_frame_interval != post_start_frame_interval:
                    current_frame_interval = post_start_frame_interval

            for future in as_completed(futures):
                text, timestamp, image_filenames, frame = future.result()
                if text:
                    if not first_phrase_found:
                        if start_phrase in text:
                            first_phrase_found = True
                            start_timestamp = timestamp
                            current_text = text
                            extracted_content.append((text, timestamp, image_filenames))
                            saved_images.extend(image_filenames)
                            previous_text = text
                            previous_images = image_filenames
                            previous_frame = frame.copy()
                    elif end_phrase in text:
                        second_phrase_found = True
                        end_timestamp = timestamp
                        extracted_content.append((text, timestamp, image_filenames))
                        saved_images.extend(image_filenames)
                        break
                    elif first_phrase_found:
                        if previous_frame is not None:
                            image_change = calculate_image_change(previous_frame, frame)
                            text_similarity_score = text_similarity(previous_text, text) if previous_text and text else 0.0

                            if image_change > MIN_CHANGE_THRESHOLD or text_similarity_score < (1 - MIN_CHANGE_THRESHOLD):
                                extracted_content.append((text, timestamp, image_filenames))
                                saved_images.extend(image_filenames)
                                previous_text = text
                                previous_images = image_filenames
                                previous_frame = frame.copy()
                        else:
                            extracted_content.append((text, timestamp, image_filenames))
                            saved_images.extend(image_filenames)
                            previous_text = text
                            previous_images = image_filenames
                            previous_frame = frame.copy()

            cap.release()

        if first_phrase_found and second_phrase_found:
            filtered_content = []
            for text, timestamp, image_filenames in extracted_content:
                if start_timestamp <= timestamp <= end_timestamp:
                    filtered_content.append((text, timestamp, image_filenames))

            print(f"DEBUG: filtered_content before Gemini: {filtered_content}")
            gemini_explanation = get_gemini_explanation(filtered_content)
            print(f"DEBUG: gemini_explanation after Gemini call: {gemini_explanation}")
            doc = create_enhanced_document(filtered_content, gemini_explanation)
            temp_docx_path = "temp_lecture_analysis.docx"

            with open(temp_docx_path, "wb") as f:
                doc.save(f)

            with open(temp_docx_path, "rb") as f:
                docx_data = io.BytesIO(f.read())
            os.remove(temp_docx_path)
            temp_docx_path = None

            all_saved_images = set()
            for _, _, image_filenames_list in filtered_content:
                all_saved_images.update(image_filenames_list)

            for filename in os.listdir(IMAGE_OUTPUT_DIR):
                file_path = os.path.join(IMAGE_OUTPUT_DIR, filename)
                if file_path not in all_saved_images:
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"Error removing unused image {file_path}: {str(e)}")
            return send_file(docx_data,
                             download_name="lecture_analysis.docx",
                             as_attachment=True)
        else:
            return jsonify({'error': 'Start or end phrase not found in the video'}), 404

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return jsonify({'error': str(e)}), 500

    finally:
        if temp_video_path and os.path.exists(temp_video_path):
            try:
                os.remove(temp_video_path)
            except Exception as e:
                app.logger.error(f"Error removing temp_video_path: {e}")

if __name__ == '__main__':
    app.run(debug=True)
