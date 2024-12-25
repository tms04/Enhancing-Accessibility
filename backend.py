import cv2
import pytesseract
import time
from langdetect import detect
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
from docx import Document
from docx.shared import Inches
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from scipy.spatial.distance import cosine
from groq import Groq
import io
from PIL import Image

# Configuration
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
VIDEO_FILE = 'AT Lecture - 1.mp4'
OUTPUT_FILE = "output_with_timestamps.txt"
OUTPUT_WORD_FILE = "lecture_analysis.docx"
IMAGE_OUTPUT_DIR = "extracted_images"
INTERVAL_MINUTES = 3
FIRST_STOP_PHRASE = "Are these Computers?"
SECOND_STOP_PHRASE = "What is a Computer?"
MAX_WORKERS = 4
TEXT_SIMILARITY_THRESHOLD = 0.7
IMAGE_SIMILARITY_THRESHOLD = 0.95

def perform_ocr(frame, timestamp):
    """Perform OCR on a video frame."""
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray)
    return text, timestamp

def filter_text(text):
    """Filter text based on language and content."""
    try:
        lang = detect(text)
        return lang == 'en' and len(text.strip()) > 0
    except:
        return False

def write_output_file(text, timestamp):
    """Write extracted text to output file."""
    with open(OUTPUT_FILE, "a") as file:
        file.write(f"{timestamp}\n{'---' * 10}\n{text}\n")

def detect_and_save_images(frame, timestamp):
    """Detect and save images from a video frame."""
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    cnts = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]

    image_filenames = []
    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        if w > 100 and h > 100:
            roi = frame[y:y+h, x:x+w]
            if is_likely_image(roi):
                if not os.path.exists(IMAGE_OUTPUT_DIR):
                    os.makedirs(IMAGE_OUTPUT_DIR)
                filename = f"{IMAGE_OUTPUT_DIR}/image_{timestamp.replace(':', '-')}_{len(image_filenames)}.png"
                cv2.imwrite(filename, roi)
                image_filenames.append(filename)
    
    return image_filenames

def is_likely_image(roi):
    """Check if an ROI is likely to be an image."""
    gray_roi = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    std_dev = np.std(gray_roi)
    return std_dev > 20

def process_frame(frame, current_time):
    """Process a single video frame."""
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(current_time))
    text, _ = perform_ocr(frame, timestamp)
    image_filenames = []
    if filter_text(text):
        write_output_file(text, timestamp)
        image_filenames = detect_and_save_images(frame, timestamp)
    return text, timestamp, image_filenames

def text_similarity(text1, text2):
    """Calculate similarity between two text segments."""
    vectorizer = TfidfVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    return 1 - cosine(vectors[0], vectors[1])

def image_similarity(img1_path, img2_path):
    """Calculate similarity between two images."""
    img1 = cv2.imread(img1_path, cv2.IMREAD_GRAYSCALE)
    img2 = cv2.imread(img2_path, cv2.IMREAD_GRAYSCALE)
    
    img1 = cv2.resize(img1, (300, 300))
    img2 = cv2.resize(img2, (300, 300))
    
    similarity = cv2.matchTemplate(img1, img2, cv2.TM_CCOEFF_NORMED)
    return np.max(similarity)

def describe_image(image_path):
    """Generate a description of the image."""
    try:
        img = cv2.imread(image_path)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        ocr_text = pytesseract.image_to_string(gray)
        
        height, width = img.shape[:2]
        brightness = np.mean(gray)
        contrast = np.std(gray)
        
        edges = cv2.Canny(gray, 100, 200)
        edge_density = np.sum(edges > 0) / (height * width)
        
        image_type = "text-heavy" if len(ocr_text.strip()) > 100 else \
                    "diagram" if edge_density > 0.1 else \
                    "photo"
        
        description = f"""Image Description:
        - Type: {image_type}
        - Dimensions: {width}x{height}
        - Content: {ocr_text.strip() if ocr_text.strip() else 'No text detected'}
        - Visual characteristics: {'High contrast with detailed edges' if edge_density > 0.1 else 'Smooth transitions with natural appearance'}
        """
        
        return description
    except Exception as e:
        print(f"Error describing image {image_path}: {str(e)}")
        return f"[Error analyzing image: {image_path}]"

def get_groq_explanation(extracted_content):
    """Get detailed explanation from Groq."""
    try:
        formatted_content = ""
        for text, timestamp, image_filenames in extracted_content:
            formatted_content += f"\n=== Timestamp: {timestamp} ===\n"
            formatted_content += f"Text Content:\n{text}\n"
            
            if image_filenames:
                formatted_content += "\nVisual Content:\n"
                for img_file in image_filenames:
                    image_desc = describe_image(img_file)
                    formatted_content += f"{image_desc}\n"
            
            formatted_content += "\n" + "="*50 + "\n"

        client = Groq(api_key="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")
        
        prompt = f"""Analyze the following lecture content, including both text and visual elements. Please provide:

        1. Comprehensive Analysis:
           - Main concepts or each point explained in detail
           - Integration of visual elements with textual content
           - Technical terminology explained
        
        2. Visual Content Analysis:
           - How the images support the lecture content
           - Key visual elements and their significance
           - Relationship between text and visual aids
           -Also identify the different images in that image if any
        
        3. Summary and Practical Applications:
           - Key takeaways
           - Real-world applications
           - Learning objectives achieved
        5.Alternate explanation of the slides.
        4.If you find any numerical in the text provided by me also solve that numerical
        Lecture Content:
        {formatted_content}

        Please structure your response with clear sections and subsections, relating the visual elements to the concepts discussed."""

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama3-8b-8192",
        )
        
        return chat_completion.choices[0].message.content
    except Exception as e:
        print(f"Error getting Groq explanation: {str(e)}")
        return None

def create_enhanced_document(extracted_content, groq_explanation):
    """Create a Word document with all content."""
    doc = Document()
    
    doc.add_heading('Lecture Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading('Original Lecture Content', 1)
    for text, timestamp, image_filenames in extracted_content:
        doc.add_heading(f"Timestamp: {timestamp}", 2)
        doc.add_paragraph(text)
        
        if image_filenames:
            doc.add_heading("Visual Content", 3)
            for image_filename in image_filenames:
                doc.add_picture(image_filename, width=Inches(6))
                desc = describe_image(image_filename)
                doc.add_paragraph(desc, style='Caption')
                
        doc.add_paragraph()
    
    if groq_explanation:
        doc.add_heading('Comprehensive Analysis', 1)
        doc.add_paragraph(groq_explanation)
    
    doc.add_heading('Technical Appendix', 1)
    doc.add_paragraph(f"""
    Processing Details:
    - Video File: {os.path.basename(VIDEO_FILE)}
    - Content Extraction Interval: {INTERVAL_MINUTES} minutes
    - Number of Images Processed: {sum(len(img) for _, _, img in extracted_content)}
    - Analysis Model: llama3-8b-8192
    """)
    
    save_attempts = 0
    while save_attempts < 3:
        try:
            if save_attempts == 0:
                output_file = OUTPUT_WORD_FILE
            else:
                output_file = f"lecture_analysis_{save_attempts}.docx"
            
            doc.save(output_file)
            print(f"Enhanced document saved as {output_file}")
            break
        except PermissionError:
            save_attempts += 1
            print(f"Permission denied when trying to save as {output_file}. Trying alternative filename...")
        except Exception as e:
            print(f"An error occurred while saving the document: {str(e)}")
            break

def main():
    """Main function to process video and create analysis."""
    cap = cv2.VideoCapture(VIDEO_FILE)
    if not cap.isOpened():
        print("Error: Could not open video file.")
        return

    start_time = time.time()
    print("Code execution started at:", time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(start_time)))

    first_phrase_found = False
    second_phrase_found = False
    extracted_content = []
    current_text = ""
    saved_images = []
    start_timestamp = None
    end_timestamp = None

    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = []
        for frame_number in range(0, frame_count, int(fps * INTERVAL_MINUTES * 60)):
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
            ret, frame = cap.read()
            if not ret:
                break
            current_time = frame_number / fps
            futures.append(executor.submit(process_frame, frame, current_time))

        for future in as_completed(futures):
            text, timestamp, image_filenames = future.result()
            if text:
                if not first_phrase_found:
                    if FIRST_STOP_PHRASE in text:
                        print(f"First stop phrase '{FIRST_STOP_PHRASE}' found at {timestamp}.")
                        first_phrase_found = True
                        start_timestamp = timestamp
                        current_text = text
                        extracted_content.append((text, timestamp, image_filenames))
                        saved_images.extend(image_filenames)
                elif SECOND_STOP_PHRASE in text:
                    print(f"Second stop phrase '{SECOND_STOP_PHRASE}' found at {timestamp}. Stopping processing.")
                    second_phrase_found = True
                    end_timestamp = timestamp
                    extracted_content.append((text, timestamp, image_filenames))
                    saved_images.extend(image_filenames)
                    break
                elif first_phrase_found and text_similarity(current_text, text) < TEXT_SIMILARITY_THRESHOLD:
                    current_text = text
                    new_images = []
                    for img in image_filenames:
                        if not any(image_similarity(img, saved_img) > IMAGE_SIMILARITY_THRESHOLD for saved_img in saved_images):
                            new_images.append(img)
                    if new_images:
                        extracted_content.append((text, timestamp, new_images))
                        saved_images.extend(new_images)

    cap.release()

    if first_phrase_found and second_phrase_found:
        filtered_content = []
        
        for text, timestamp, image_filenames in extracted_content:
            if start_timestamp <= timestamp <= end_timestamp:
                filtered_content.append((text, timestamp, image_filenames))
        
        groq_explanation = get_groq_explanation(filtered_content)
        create_enhanced_document(filtered_content, groq_explanation)
        
        all_saved_images = set()
        for _, _, image_filenames in filtered_content:
            all_saved_images.update(image_filenames)
        
        for filename in os.listdir(IMAGE_OUTPUT_DIR):
            file_path = os.path.join(IMAGE_OUTPUT_DIR, filename)
            if file_path not in all_saved_images:
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Error removing unused image {file_path}: {str(e)}")

    else:
        print("Error: Start or end phrase not found in the video.")

    end_time = time.time()
    print("Code execution ended at:", time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(end_time)))
    print("Total time taken to execute:", end_time - start_time)

if __name__ == "__main__":
    main()